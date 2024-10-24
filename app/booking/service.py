"""Service for booking domain with business logic."""
import time
from datetime import datetime
from typing import TypeAlias

from docx import Document
from fastapi import HTTPException

from app.booking.model import Booking
from app.booking.repository import BookingRepository
from app.config import BASE_DIR


PathToFile: TypeAlias = str


class BookingService:
    def __init__(self):
        self.repository = BookingRepository()

    def create_booking(self, booking) -> int:
        """Process to create booking from user's input data."""
        if booking.check_in_date >= booking.check_out_date:
            raise HTTPException(status_code=400,
                                detail="Дата начала бронирования должна быть больше даты окончания")
        occupied_rooms_count = self.repository.get_occupied_rooms_count(booking.room_number,
                                                                        booking.check_in_date,
                                                                        booking.check_out_date)
        if occupied_rooms_count > 0:
            raise HTTPException(status_code=400,
                                detail="Переговорная комната недоступна в указанное время")
        created_id = self.repository.create_new_booking(booking)
        return created_id

    def get_booking_by_id(self, booking_id: int) -> Booking:
        """Get Booking entity by id."""
        booking = self.repository.get_booking(booking_id)
        if not booking:
            raise HTTPException(status_code=404,
                                detail=f"Бронирование с идентификатором {booking_id} не найдено.")
        return booking

    def get_all_bookings(self, filters) -> list[Booking]:
        """All filtered bookings."""
        all_bookings = self.repository.get_all_bookings(filters)
        for booking in all_bookings:
            booking.is_empty_now = None
            if filters.get('is_current_day') is True:
                if booking.check_in_date <= datetime.now() <= booking.check_out_date:
                    booking.is_empty_now = True
                else:
                    booking.is_empty_now = False
        return all_bookings

    def get_data_for_report(self, filters: dict) -> list[Booking]:
        return self.repository.get_report_data(filters)

    @staticmethod
    def create_report(bookings: list[Booking], filters) -> PathToFile:
        """Create table structure with input data."""
        document = Document()
        if filters.get('room_number') is not None:
            heading = f"Список бронирований комнаты номер - {filters.get('room_number')}"

            document.add_heading(heading, level=0)

            table = document.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            table_cells = table.rows[0].cells
            table_cells[0].text = "Кто забронировал"
            table_cells[1].text = "Время создания бронирования"
            table_cells[2].text = "Цель"

            for booking in bookings:
                row_cells = table.add_row().cells
                row_cells[0].text = booking.who_booked
                row_cells[1].text = booking.created_at.strftime('%X %d-%m-%y')
                row_cells[2].text = booking.purpose or ""

        else:
            heading = "Список бронирований комнат"

            document.add_heading(heading, level=0)

            table = document.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            table_cells = table.rows[0].cells

            table_cells[0].text = "Номер комнаты"
            table_cells[1].text = "Кто забронировал"
            table_cells[2].text = "Время создания бронирования"
            table_cells[3].text = "Цель"
            bookings = sorted(bookings, key=lambda element: element.room_number)
            for booking in bookings:
                row_cells = table.add_row().cells
                row_cells[0].text = str(booking.room_number)
                row_cells[1].text = booking.who_booked
                row_cells[2].text = booking.created_at.strftime('%X %d-%m-%y')
                row_cells[3].text = booking.purpose or ""

        report_path = str(BASE_DIR) + f"/storage/reports/bookings_report{time.time_ns()}.docx"
        document.save(report_path)
        return report_path
